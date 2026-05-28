from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margens ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border_bottom(cell, color='E2E8F0'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def add_heading(doc, number, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run_num = p.add_run(f'{number}. ')
    run_num.bold      = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    run_txt = p.add_run(text)
    run_txt.bold      = True
    run_txt.font.size = Pt(14)
    run_txt.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

def add_body_rich(doc, parts):
    """parts: list of (text, bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold       = bold
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

def add_bullet(doc, text, bold_part=None):
    p  = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_part:
        idx = text.find(bold_part)
        if idx >= 0:
            before = text[:idx]
            after  = text[idx+len(bold_part):]
            if before:
                r = p.add_run(before); r.font.size = Pt(11)
            rb = p.add_run(bold_part); rb.bold = True; rb.font.size = Pt(11)
            if after:
                r = p.add_run(after); r.font.size = Pt(11)
            return
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_numbered(doc, text):
    p   = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_table(doc, headers, rows, footer_row=None):
    cols  = len(headers)
    table = doc.add_table(rows=1 + len(rows) + (1 if footer_row else 0), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], '1E3A5F')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # linhas
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = 'F8FAFC' if r_idx % 2 == 0 else 'FFFFFF'
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            set_cell_bg(row_cells[i], bg)
            set_cell_border_bottom(row_cells[i])
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size      = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # rodapé opcional
    if footer_row:
        foot_cells = table.rows[-1].cells
        for i, cell_text in enumerate(footer_row):
            foot_cells[i].text = cell_text
            set_cell_bg(foot_cells[i], '1E3A5F')
            for para in foot_cells[i].paragraphs:
                for run in para.runs:
                    run.bold           = True
                    run.font.size      = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

def add_alert(doc, icon, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f'{icon} {title}  ')
    r1.bold           = True
    r1.font.size      = Pt(11)
    r1.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    r2 = p.add_run(text)
    r2.font.size      = Pt(11)
    r2.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)

# ══════════════════════════════════════════════════════════
# CAPA / CABEÇALHO
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('RELATÓRIO TÉCNICO')
run.bold           = True
run.font.size      = Pt(10)
run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Uso de Chave API para Conversão Automatizada\nde Fórmulas LaTeX no Canvas LMS para UDAS')
run.bold           = True
run.font.size      = Pt(20)
run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

doc.add_paragraph()

meta_table = doc.add_table(rows=1, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = meta_table.rows[0].cells
meta = [('Equipe', 'Gestão do Ambiente Virtual'),
        ('Data',   'Maio de 2026'),
        ('Versão', '1.0'),
        ('Classif.', 'Interno')]
for i, (label, value) in enumerate(meta):
    p_label = cells[i].paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_label.add_run(label + '\n')
    rl.font.size      = Pt(8)
    rl.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    rv = p_label.add_run(value)
    rv.bold           = True
    rv.font.size      = Pt(10)
    rv.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ══════════════════════════════════════════════════════════
add_heading(doc, 1, 'Introdução')
add_body(doc, 'O Canvas LMS (Instructure) é a plataforma de ensino adotada pela instituição para disponibilização de conteúdos educacionais. Utilizamos os materiais disponibilizados nas FICs para incluir o conteúdo das UDAs no Canvas. Nesse processo, é frequente a presença de fórmulas matemáticas, equações e expressões científicas nas FICs referentes às UDAs, que precisam ser inseridas digitalmente de forma correta e acessível.')
add_body_rich(doc, [
    ('A plataforma Canvas oferece suporte nativo ao formato ', False),
    ('LaTeX', True),
    (', padrão amplamente utilizado para representação de fórmulas matemáticas em ambientes digitais. Quando uma fórmula é inserida via LaTeX, ela é renderizada visualmente pela plataforma e também fica legível por tecnologias assistivas, como leitores de tela — essencial para garantir a ', False),
    ('acessibilidade do conteúdo', True),
    (' a estudantes com deficiência visual.', False),
])
add_body(doc, 'Este relatório justifica a adoção de uma chave de API de inteligência artificial (OpenAI ou OpenRouter) para viabilizar o uso de uma extensão do Chrome com o objetivo de automatizar a conversão de imagens de fórmulas para código LaTeX.')
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 2. PROBLEMA ATUAL
# ══════════════════════════════════════════════════════════
add_heading(doc, 2, 'O Problema Atual')

add_subheading(doc, 'Origem das fórmulas nos materiais FIC')
add_body(doc, 'Os materiais recebidos em formato FIC frequentemente contêm fórmulas matemáticas apresentadas como imagens — arquivos digitalizados ou elementos visuais incorporados em documentos. Essas imagens não possuem nenhuma informação textual subjacente.')

add_subheading(doc, 'Por que imagens não são suficientes')
add_alert(doc, '♿', 'Problema de Acessibilidade',
    'Leitores de tela utilizados por pessoas com deficiência visual não conseguem interpretar o conteúdo de uma imagem. Uma fórmula inserida como imagem simplesmente não existe para esse grupo de estudantes — tornando o conteúdo inacessível e potencialmente em desconformidade com as diretrizes WCAG 2.1.')

add_subheading(doc, 'O fluxo atual sem automação')
add_numbered(doc, 'Identificar a fórmula no material FIC (imagem)')
add_numbered(doc, 'Abrir o editor LaTeX no Canvas')
add_numbered(doc, 'Tentar reconstituir manualmente o código LaTeX correspondente à imagem')
add_numbered(doc, 'Quando o conhecimento técnico não é suficiente, copiar a imagem e colar no ChatGPT gratuito para obter sugestão de código')
add_numbered(doc, 'Copiar o código sugerido e colar no editor do Canvas')
add_numbered(doc, 'Verificar se a renderização está correta; ajustar se necessário')
doc.add_paragraph()

add_subheading(doc, 'Limitações graves do fluxo atual')
add_table(doc,
    ['Problema', 'Impacto'],
    [
        ['Reconstituição manual exige conhecimento avançado em LaTeX', 'Nem todos os membros da equipe dominam a sintaxe'],
        ['Processo lento por fórmula', 'Cada fórmula pode levar de 3 a 10 minutos'],
        ['ChatGPT gratuito atinge limite de uso', 'Bloqueia novas solicitações no meio do trabalho'],
        ['Dependência de ferramentas externas não integradas', 'Múltiplos passos manuais entre abas e aplicações'],
        ['Sujeito a erro humano', 'Código LaTeX incorreto gera fórmula errada sem aviso imediato'],
    ]
)
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 3. SOLUÇÃO PROPOSTA
# ══════════════════════════════════════════════════════════
add_heading(doc, 3, 'Solução Proposta')
add_body_rich(doc, [
    ('A solução proposta consiste na utilização de uma extensão do Chrome com uma funcionalidade específica para este problema: a ', False),
    ('conversão automática de imagem de fórmula para código LaTeX com um único clique', True),
    ('. Para que isso seja possível, é necessária uma chave de API de inteligência artificial que permita à extensão processar as imagens das fórmulas.', False),
])

add_subheading(doc, 'Como funciona')
add_numbered(doc, 'O colaborador identifica a imagem da fórmula no material (página web, PDF no navegador ou material do Canvas)')
add_numbered(doc, 'Clica com o botão direito sobre a imagem')
add_numbered(doc, 'Seleciona "Converter fórmula para LaTeX" no menu de contexto')
add_numbered(doc, 'A extensão envia a imagem automaticamente para a API de IA configurada')
add_numbered(doc, 'O modelo analisa a imagem e identifica a fórmula matemática')
add_numbered(doc, 'O código LaTeX é copiado diretamente para a área de transferência — Ctrl+V já está pronto')
add_numbered(doc, 'No editor LaTeX do Canvas, basta colar — a fórmula aparece renderizada corretamente')
doc.add_paragraph()
add_alert(doc, '⚡', 'Tempo total do processo: aproximadamente 5 a 10 segundos.',
    'Sem abrir nenhuma aba adicional, sem copiar para outro site, sem aguardar interfaces carregarem.')

add_subheading(doc, 'Tecnologia utilizada')
add_body_rich(doc, [
    ('A extensão utiliza o modelo ', False),
    ('GPT-4o (OpenAI)', True),
    (', que possui capacidade de visão computacional — analisa imagens e interpreta seu conteúdo. É reconhecido como estado da arte para interpretação visual e conversão de conteúdo matemático.', False),
])
add_table(doc,
    ['Tipo de chave', 'Prefixo', 'Endpoint utilizado'],
    [
        ['OpenAI (direta)', 'sk-…', 'api.openai.com'],
        ['OpenRouter',      'sk-or-…', 'openrouter.ai'],
    ]
)
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 4. BENEFÍCIOS
# ══════════════════════════════════════════════════════════
add_heading(doc, 4, 'Benefícios')
add_bullet(doc, 'Produtividade: redução de 80 a 95% no tempo por fórmula. Em um material com 10 fórmulas, economia de 40 a 120 minutos.', 'Produtividade:')
add_bullet(doc, 'Acessibilidade garantida: toda fórmula em LaTeX é legível por leitores de tela, cumprindo as diretrizes WCAG 2.1 nível AA.', 'Acessibilidade garantida:')
add_bullet(doc, 'Sem limite de uso: diferente do ChatGPT gratuito, a API é ilimitada — cada requisição é cobrada individualmente, sem teto de conversas.', 'Sem limite de uso:')
add_bullet(doc, 'Alta precisão: GPT-4o interpreta frações, somatórios, integrais, vetores e expressões complexas com alta fidelidade.', 'Alta precisão:')
add_bullet(doc, 'Integrado ao fluxo: funciona direto no navegador, sem trocar de aba ou abrir ferramentas externas.', 'Integrado ao fluxo:')
add_bullet(doc, 'Democrático: não exige conhecimento em LaTeX — qualquer membro da equipe consegue usar com um clique.', 'Democrático:')
doc.add_paragraph()

add_subheading(doc, 'Comparativo de tempo estimado')
add_table(doc,
    ['Etapa', 'Antes', 'Depois'],
    [
        ['Identificar e copiar fórmula',  '~1 min',       '5 seg'],
        ['Obter código LaTeX',            '3 – 10 min',   'automático'],
        ['Colar e verificar no Canvas',   '1 – 2 min',    '~30 seg'],
    ],
    footer_row=['Total estimado por fórmula', '5 – 13 minutos', '~1 minuto']
)
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 5. JUSTIFICATIVA TÉCNICA
# ══════════════════════════════════════════════════════════
add_heading(doc, 5, 'Justificativa Técnica da Escolha')
add_table(doc,
    ['Critério', 'ChatGPT Gratuito', 'API (OpenAI / OpenRouter)'],
    [
        ['Limite de uso',                   'Sim — interrompível',  'Não — pay-as-you-go'],
        ['Integração com extensão',         'Impossível',           'Sim, via REST'],
        ['Automação (1 clique)',             'Impossível',           'Sim'],
        ['Controle da chave',               'Não',                  'Sim — revogável a qualquer momento'],
        ['Interpretação de imagem',         'Manual (usuário descreve)', 'Automática (visão computacional)'],
        ['Custo',                           'Gratuito com limitações', 'Muito baixo — centavos por fórmula'],
    ]
)
add_body_rich(doc, [
    ('O GPT-4o é um modelo ', False),
    ('multimodal', True),
    (' — processa tanto texto quanto imagens. Sua capacidade de interpretar fórmulas matemáticas a partir de imagens é vastamente superior a modelos menores ou soluções OCR tradicionais, que frequentemente falham com notações complexas.', False),
])
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 6. CUSTOS E VIABILIDADE
# ══════════════════════════════════════════════════════════
add_heading(doc, 6, 'Custos e Viabilidade')
add_body(doc, 'A API da OpenAI utiliza o modelo pay-as-you-go: paga-se apenas pelo que for usado, por token processado. Não há mensalidade, não há contrato mínimo.')

add_subheading(doc, 'Estimativa de custo por conversão')
add_table(doc,
    ['Item', 'Quantidade estimada', 'Custo (USD)'],
    [
        ['Tokens de entrada (imagem + prompt)', '~200 – 500 tokens', '~$0,0010'],
        ['Tokens de saída (código LaTeX)',       '~50 – 200 tokens',  '~$0,0010'],
    ],
    footer_row=['Total por fórmula', '', '~$0,001 a $0,003']
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('Custo médio por fórmula: $0,002 USD ≈ R$ 0,01 — menos de 1 centavo')
run.bold           = True
run.font.size      = Pt(13)
run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)

add_subheading(doc, 'Projeção de uso mensal')
add_table(doc,
    ['Cenário', 'Fórmulas / mês', 'Custo est. (USD)', 'Custo est. (BRL)'],
    [
        ['Uso leve (1–2 FICs)',             '~50',    '~$0,10', '~R$ 0,57'],
        ['Uso moderado (5–8 FICs)',         '~300',   '~$0,60', '~R$ 3,42'],
        ['Uso intenso (FICs + revisões)',   '~1.000', '~$2,00', '~R$ 11,40'],
    ]
)
add_alert(doc, '💡', 'Recomendação:',
    'Definir um limite de gasto mensal de $5,00 USD na conta da OpenAI. Esse valor cobre mais de 2.000 conversões — suficiente para qualquer cenário de uso real — e evita cobranças inesperadas.')
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 7. PRECAUÇÕES E SEGURANÇA
# ══════════════════════════════════════════════════════════
add_heading(doc, 7, 'Precauções e Segurança')

add_subheading(doc, 'Sobre a chave de API')
add_bullet(doc, 'A chave é armazenada localmente no navegador do colaborador, dentro do armazenamento seguro da extensão Chrome — não é enviada para nenhum servidor da equipe')
add_bullet(doc, 'A requisição vai diretamente do navegador para o endpoint da OpenAI/OpenRouter, sem intermediários')
add_bullet(doc, 'A chave pode ser revogada a qualquer momento pelo painel da OpenAI, sem impacto a outros sistemas')
add_bullet(doc, 'Recomenda-se definir um limite de gasto mensal na conta da OpenAI para evitar cobranças inesperadas')
add_bullet(doc, 'A chave deve ser gerada em uma conta vinculada ao responsável técnico ou à instituição')
doc.add_paragraph()

add_subheading(doc, 'Sobre os dados enviados')
add_bullet(doc, 'A extensão envia apenas a imagem da fórmula para a API — nenhum dado da página, do usuário ou do conteúdo do Canvas é transmitido')
add_bullet(doc, 'Os dados de API não são usados para treinar modelos, conforme política da OpenAI para clientes de API (diferente do ChatGPT gratuito)')
add_bullet(doc, 'Nenhum dado de identificação pessoal é transmitido nas requisições')
doc.add_paragraph()

add_subheading(doc, 'Boas práticas recomendadas')
add_bullet(doc, 'Não compartilhar a chave de API em canais públicos, e-mails ou documentos não protegidos')
add_bullet(doc, 'Monitorar o uso pelo painel da OpenAI ao menos uma vez por mês')
add_bullet(doc, 'Revogar e substituir a chave imediatamente em caso de suspeita de vazamento')
add_bullet(doc, 'Manter a extensão sempre atualizada para garantir as últimas correções de segurança')
doc.add_paragraph()
add_alert(doc, '⚠️', 'Atenção:',
    'A chave de API possui valor financeiro — quem a tiver pode realizar requisições cobradas na conta. Trate-a como uma senha: não compartilhe, não registre em documentos públicos, e revogue imediatamente caso haja suspeita de comprometimento.')
add_hr(doc)

# ══════════════════════════════════════════════════════════
# 8. CONCLUSÃO
# ══════════════════════════════════════════════════════════
add_heading(doc, 8, 'Conclusão')
add_body(doc, 'A adoção da chave de API para conversão automática de fórmulas LaTeX resolve um gargalo real e recorrente no fluxo de produção de materiais no Canvas LMS.')
add_body(doc, 'O processo atual é lento, dependente de conhecimento técnico especializado e sujeito a interrupções impostas por limites de ferramentas gratuitas. A solução proposta — uma extensão do Chrome integrada à API — reduz o tempo de inserção de cada fórmula de vários minutos para menos de 10 segundos, garante acessibilidade digital ao conteúdo e elimina a dependência de ferramentas externas.')
add_body_rich(doc, [
    ('O custo operacional é ', False),
    ('negligenciável', True),
    (' (inferior a R$ 0,02 por fórmula). Em um único dia de trabalho com múltiplos materiais FIC, a solução automatizada economiza horas de trabalho manual a um custo de centavos.', False),
])
add_body_rich(doc, [
    ('A relação custo-benefício é ', False),
    ('amplamente favorável à implementação imediata.', True),
])

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento interno — Equipe de Gestão do Ambiente Virtual  ·  Maio de 2026')
run.font.size      = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save(r'c:\Users\Casa\servicetime-v2\relatorio-api-latex.docx')
print('Arquivo gerado com sucesso.')
