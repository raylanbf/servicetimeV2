"""
Gera relatorio-api-latex.docx
Abordagem limpa: sem tabelas aninhadas, sem cantSplit em tabelas grandes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Página A4 com margens padrão ─────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── XML helpers ───────────────────────────────────────────────────────────────
def rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def shd(el_tc_or_pPr, fill, is_para=False):
    """Aplica cor de fundo numa célula ou parágrafo."""
    owner = el_tc_or_pPr
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill.lstrip('#'))
    owner.append(s)

def cell_shd(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:shd')): tcPr.remove(o)
    shd(tcPr, fill)

def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcMar')): tcPr.remove(o)
    m = OxmlElement('w:tcMar')
    for side, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)

def cell_borders_none(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcBorders')): tcPr.remove(o)
    b = OxmlElement('w:tcBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'none'); b.append(e)
    tcPr.append(b)

def cell_border_left(cell, color='2563EB', sz='20'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcBorders')): tcPr.remove(o)
    b = OxmlElement('w:tcBorders')
    for s in ('top','right','bottom','insideH','insideV'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'none'); b.append(e)
    lft = OxmlElement('w:left')
    lft.set(qn('w:val'),'single'); lft.set(qn('w:sz'),sz)
    lft.set(qn('w:space'),'0'); lft.set(qn('w:color'),color.lstrip('#'))
    b.append(lft); tcPr.append(b)

def cell_all_borders(cell, color='E2E8F0', sz='6'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcBorders')): tcPr.remove(o)
    b = OxmlElement('w:tcBorders')
    for s in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color.lstrip('#'))
        b.append(e)
    tcPr.append(b)

def cell_bottom_border(cell, color='E2E8F0', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn('w:tcBorders')): tcPr.remove(o)
    b = OxmlElement('w:tcBorders')
    for s in ('top','left','right','insideH','insideV'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'none'); b.append(e)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),sz)
    bot.set(qn('w:space'),'0'); bot.set(qn('w:color'),color.lstrip('#'))
    b.append(bot); tcPr.append(b)

def tbl_width_pct(table, pct=5000):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    for o in tblPr.findall(qn('w:tblW')): tblPr.remove(o)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(pct)); w.set(qn('w:type'), 'pct')
    tblPr.append(w)

def tbl_no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    for o in tblPr.findall(qn('w:tblBorders')): tblPr.remove(o)
    b = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'none'); b.append(e)
    tblPr.append(b)
    tbl_width_pct(table)

def para_spc(p, bef=0, aft=0, line=None):
    pPr = p._p.get_or_add_pPr()
    for o in pPr.findall(qn('w:spacing')): pPr.remove(o)
    s = OxmlElement('w:spacing')
    s.set(qn('w:before'), str(bef)); s.set(qn('w:after'), str(aft))
    if line:
        s.set(qn('w:line'), str(int(line*20))); s.set(qn('w:lineRule'),'exact')
    pPr.append(s)

def para_bg(p, fill):
    pPr = p._p.get_or_add_pPr()
    for o in pPr.findall(qn('w:shd')): pPr.remove(o)
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'), fill.lstrip('#'))
    pPr.append(s)

def para_indent(p, left_cm=0, right_cm=0):
    p.paragraph_format.left_indent  = Cm(left_cm)
    p.paragraph_format.right_indent = Cm(right_cm)

def run(p, text, bold=False, size=11, color='374151', italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    return r

def hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bd  = OxmlElement('w:pBdr')
    b   = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'CBD5E1')
    bd.append(b); pPr.append(bd)
    para_spc(p, 180, 180)

def sp(doc, pt=6):
    p = doc.add_paragraph()
    para_spc(p, 0, int(pt*20))

# ═══════════════════════════════════════════════════════════════════════════
# CABEÇALHO — duas tabelas full-bleed (título + metadados)
# ═══════════════════════════════════════════════════════════════════════════
TWIPS_PER_CM = 567
MARGIN_CM    = 2.5
A4_TWIPS     = int(21 * TWIPS_PER_CM)
INDENT_TWIPS = int(MARGIN_CM * TWIPS_PER_CM)
INNER_PAD    = INDENT_TWIPS + 80  # padding interno alinhado com margens do corpo

def apply_full_bleed(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tbl_no_borders(table)
    for o in tblPr.findall(qn('w:tblW')): tblPr.remove(o)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(A4_TWIPS)); w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    for o in tblPr.findall(qn('w:tblInd')): tblPr.remove(o)
    ind = OxmlElement('w:tblInd')
    ind.set(qn('w:w'), str(-INDENT_TWIPS)); ind.set(qn('w:type'), 'dxa')
    tblPr.append(ind)

def add_line_break(para):
    """Insere <w:br/> real dentro do último run do parágrafo."""
    br = OxmlElement('w:br')
    para.runs[-1]._r.append(br)

# ── Linha de título ───────────────────────────────────────────────────────
hdr1 = doc.add_table(rows=1, cols=1)
apply_full_bleed(hdr1)
c0 = hdr1.rows[0].cells[0]
cell_shd(c0, '1E3A5F')
cell_borders_none(c0)
cell_margins(c0, top=260, bottom=60, left=INNER_PAD, right=INNER_PAD)
p0 = c0.paragraphs[0]
para_spc(p0, 0, 40)
run(p0, 'RELATÓRIO TÉCNICO', bold=True, size=8, color='93C5FD')
p1 = c0.add_paragraph()
para_spc(p1, 0, 0)
# quebra de linha real com <w:br/> no ponto certo
run(p1, 'Uso de Chave API para Conversão Automatizada', bold=True, size=16, color='FFFFFF')
add_line_break(p1)
run(p1, 'de Fórmulas LaTeX no Canvas LMS para UDAS',   bold=True, size=16, color='FFFFFF')

# ── Linha de metadados — larguras explícitas por coluna ──────────────────
# Col0 (Equipe): 38% | Col1 (Data): 24% | Col2 (Versão): 18% | Col3 (Classif.): 20%
COL_WIDTHS_CM = [8.0, 5.0, 4.0, 4.0]  # soma = 21 cm = largura A4

hdr2 = doc.add_table(rows=1, cols=4)
apply_full_bleed(hdr2)
meta = [
    ('Equipe',   'Gestão do Ambiente Virtual'),
    ('Data',     'Maio de 2026'),
    ('Versão',   '1.0'),
    ('Classif.', 'Interno'),
]
for i, (lbl, val) in enumerate(meta):
    c = hdr2.rows[0].cells[i]
    c.width = Cm(COL_WIDTHS_CM[i])
    cell_shd(c, '162D4A')
    cell_borders_none(c)
    lpad = INNER_PAD if i == 0 else 140
    rpad = INNER_PAD if i == 3 else 140
    cell_margins(c, top=120, bottom=200, left=lpad, right=rpad)
    pm = c.paragraphs[0]
    para_spc(pm, 0, 0)
    run(pm, lbl, size=8, color='93C5FD')
    add_line_break(pm)
    run(pm, val, bold=True, size=10, color='FFFFFF')

sp(doc, 16)

# ── Helpers de conteúdo ──────────────────────────────────────────────────────
def section_heading(doc, num, title):
    p = doc.add_paragraph()
    para_bg(p, 'EFF6FF')
    para_spc(p, 200, 60)
    para_indent(p, 0, 0)
    # número em bloco azul simulado com negrito colorido
    rn = run(p, f'  {num}  ', bold=True, size=13, color='FFFFFF')
    rn.font.highlight_color = None
    # background no run via rPr shd
    rPr = rn._r.get_or_add_rPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'1E3A5F')
    rPr.append(s)
    run(p, f'   {title}', bold=True, size=14, color='1E3A5F')

def subh(doc, text):
    p = doc.add_paragraph()
    para_spc(p, 200, 60)
    run(p, text.upper(), bold=True, size=9, color='334155')

def body(doc, parts, aft=140):
    p = doc.add_paragraph()
    para_spc(p, 0, aft, line=16)
    for text, bold, color in parts:
        run(p, text, bold=bold, color=color)
    return p

def body_p(doc, text):
    body(doc, [(text, False, '374151')])

def numbered_list(doc, items):
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph()
        para_spc(p, 0, 60, line=15)
        para_indent(p, 0.2)
        run(p, f'{i}.  ', bold=True, size=11, color='1D4ED8')
        run(p, text, color='374151')

def bullet_list(doc, items):
    for text in items:
        p = doc.add_paragraph()
        para_spc(p, 0, 60, line=15)
        para_indent(p, 0.2)
        run(p, '✓  ', bold=True, size=11, color='16A34A')
        run(p, text, color='374151')

def alert(doc, icon, title, text, bg='EFF6FF', border='2563EB', fc='1E40AF'):
    t = doc.add_table(rows=1, cols=2)
    tbl_no_borders(t)
    tbl_width_pct(t)

    ci = t.rows[0].cells[0]
    ci.width = Cm(1.2)
    cell_shd(ci, bg)
    cell_border_left(ci, color=border)
    cell_margins(ci, top=140, bottom=140, left=100, right=80)
    pi = ci.paragraphs[0]
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spc(pi, 0, 0)
    run(pi, icon, size=14, color=fc)

    ct = t.rows[0].cells[1]
    cell_shd(ct, bg)
    cell_borders_none(ct)
    cell_margins(ct, top=140, bottom=140, left=120, right=160)
    pt = ct.paragraphs[0]
    para_spc(pt, 0, 60, line=15)
    run(pt, title + '\n', bold=True, size=11, color=fc)
    run(pt, text, size=11, color=fc)

    sp(doc, 8)

def data_table(doc, headers, rows, footer=None):
    cols  = len(headers)
    total = 1 + len(rows) + (1 if footer else 0)
    t = doc.add_table(rows=total, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_width_pct(t)
    tbl_no_borders(t)

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_shd(c, '1E3A5F')
        cell_borders_none(c)
        cell_margins(c, top=120, bottom=120, left=140, right=80)
        p = c.paragraphs[0]
        para_spc(p, 0, 0)
        run(p, h, bold=True, size=10, color='FFFFFF')

    for ri, row_data in enumerate(rows):
        bg = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            cell_shd(c, bg)
            cell_bottom_border(c)
            cell_margins(c, top=100, bottom=100, left=140, right=80)
            p = c.paragraphs[0]
            para_spc(p, 0, 0, line=15)
            run(p, val, size=10, color='334155')

    if footer:
        for ci, val in enumerate(footer):
            c = t.rows[-1].cells[ci]
            cell_shd(c, '1E3A5F')
            cell_borders_none(c)
            cell_margins(c, top=120, bottom=120, left=140, right=80)
            p = c.paragraphs[0]
            para_spc(p, 0, 0)
            run(p, val, bold=True, size=10, color='FFFFFF')

    sp(doc, 10)

def cards(doc, items):
    """Grade 2 colunas. Cada linha é uma tabela independente."""
    pairs = [items[i:i+2] for i in range(0, len(items), 2)]
    for pair in pairs:
        t = doc.add_table(rows=1, cols=2)
        tbl_no_borders(t)
        tbl_width_pct(t)
        for ci in range(2):
            c = t.rows[0].cells[ci]
            if ci < len(pair):
                icon, title, desc = pair[ci]
                cell_shd(c, 'F8FAFC')
                cell_all_borders(c)
                cell_margins(c, top=160, bottom=160, left=180, right=180)
                p = c.paragraphs[0]
                para_spc(p, 0, 50, line=16)
                run(p, icon + '\n', size=16, color='1E3A5F')
                run(p, title + '\n', bold=True, size=11, color='1E293B')
                run(p, desc, size=10, color='64748B')
            else:
                cell_shd(c, 'FFFFFF')
                cell_borders_none(c)
        sp(doc, 5)
    sp(doc, 6)

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 1, 'Introdução')
body_p(doc, 'O Canvas LMS (Instructure) é a plataforma de ensino adotada pela instituição para disponibilização de conteúdos educacionais. Utilizamos os materiais disponibilizados nas FICs para incluir o conteúdo das UDAs no Canvas. Nesse processo, é frequente a presença de fórmulas matemáticas, equações e expressões científicas nas FICs referentes às UDAs, que precisam ser inseridas digitalmente de forma correta e acessível.')
body(doc, [
    ('A plataforma Canvas oferece suporte nativo ao formato ', False, '374151'),
    ('LaTeX', True, '1E3A5F'),
    (', padrão amplamente utilizado para representação de fórmulas matemáticas em ambientes digitais. Quando uma fórmula é inserida via LaTeX, ela é renderizada visualmente e fica legível por tecnologias assistivas, como leitores de tela — essencial para garantir a ', False, '374151'),
    ('acessibilidade do conteúdo', True, '1E3A5F'),
    (' a estudantes com deficiência visual.', False, '374151'),
])
body_p(doc, 'Este relatório justifica a adoção de uma chave de API de inteligência artificial (OpenAI ou OpenRouter) para viabilizar o uso de uma extensão do Chrome com o objetivo de automatizar a conversão de imagens de fórmulas para código LaTeX.')
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 2. PROBLEMA ATUAL
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 2, 'O Problema Atual')

subh(doc, 'Origem das fórmulas nos materiais FIC')
body_p(doc, 'Os materiais recebidos em formato FIC frequentemente contêm fórmulas matemáticas apresentadas como imagens — arquivos digitalizados ou elementos visuais incorporados em documentos. Essas imagens não possuem nenhuma informação textual subjacente.')

subh(doc, 'Por que imagens não são suficientes')
alert(doc, '♿', 'Problema de Acessibilidade',
    'Leitores de tela utilizados por pessoas com deficiência visual não conseguem interpretar o conteúdo de uma imagem. Uma fórmula inserida como imagem simplesmente não existe para esse grupo de estudantes — tornando o conteúdo inacessível e potencialmente em desconformidade com as diretrizes WCAG 2.1.')

subh(doc, 'O fluxo atual sem automação')
numbered_list(doc, [
    'Identificar a fórmula no material FIC (imagem)',
    'Abrir o editor LaTeX no Canvas',
    'Tentar reconstituir manualmente o código LaTeX correspondente à imagem',
    'Quando o conhecimento técnico não é suficiente, copiar a imagem e colar no ChatGPT gratuito para obter sugestão de código',
    'Copiar o código sugerido e colar no editor do Canvas',
    'Verificar se a renderização está correta; ajustar se necessário',
])
sp(doc, 8)

subh(doc, 'Limitações graves do fluxo atual')
data_table(doc,
    ['Problema', 'Impacto'],
    [
        ['Reconstituição manual exige conhecimento avançado em LaTeX', 'Nem todos os membros da equipe dominam a sintaxe'],
        ['Processo lento por fórmula', 'Cada fórmula pode levar de 3 a 10 minutos'],
        ['ChatGPT gratuito atinge limite de uso', 'Bloqueia novas solicitações no meio do trabalho'],
        ['Dependência de ferramentas externas não integradas', 'Múltiplos passos manuais entre abas e aplicações'],
        ['Sujeito a erro humano', 'Código LaTeX incorreto gera fórmula errada sem aviso imediato'],
    ]
)
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 3. COMO OBTER ACESSO À API
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 3, 'Como Obter Acesso à API')
body_p(doc, 'Para utilizar a solução proposta, é necessário criar uma conta em uma das plataformas abaixo, adicionar créditos e gerar uma chave de API. O processo é simples e leva menos de 10 minutos.')

subh(doc, 'Opção A — OpenAI (chave direta, prefixo sk-…)')
numbered_list(doc, [
    'Acesse platform.openai.com e clique em "Sign up" para criar uma conta (ou "Log in" se já possuir)',
    'Após o login, acesse o menu "Settings" > "Billing"',
    'Clique em "Add payment method" e cadastre um cartão de crédito/débito internacional',
    'Em "Usage limits", defina o limite mensal recomendado de $5,00 USD para controle de gastos',
    'Acesse o menu "API keys" no painel lateral e clique em "Create new secret key"',
    'Dê um nome à chave (ex: "Canvas LaTeX") e clique em "Create secret key"',
    'Copie e guarde a chave gerada — ela começa com sk-… e é exibida apenas uma vez',
    'Cole a chave no campo "Chave API" da extensão Chrome em ⚙ Configurações',
])
sp(doc, 10)

subh(doc, 'Opção B — OpenRouter (prefixo sk-or-…, alternativa com mais flexibilidade)')
numbered_list(doc, [
    'Acesse openrouter.ai e clique em "Sign in" para criar uma conta (aceita login com Google)',
    'Após o login, acesse "Credits" no menu superior',
    'Clique em "Buy credits" e adicione créditos via cartão (valor mínimo geralmente de $5,00 USD)',
    'Acesse "Keys" no menu superior e clique em "Create Key"',
    'Dê um nome à chave e defina um limite de crédito opcional para segurança',
    'Copie a chave gerada — ela começa com sk-or-… e é exibida apenas uma vez',
    'Cole a chave no campo "Chave API" da extensão Chrome em ⚙ Configurações',
])
sp(doc, 8)

alert(doc, '💡', 'Qual plataforma escolher?',
    'Ambas funcionam com a extensão. A OpenAI é a opção mais direta (mesmo provedor do GPT-4o). O OpenRouter oferece mais opções de modelos e pode ser vantajoso para quem já usa outras IAs. Para o uso descrito neste relatório, qualquer uma das duas atende plenamente.')

hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 4. SOLUÇÃO PROPOSTA
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 4, 'Solução Proposta')
body(doc, [
    ('A solução proposta consiste na utilização de uma extensão do Chrome com uma funcionalidade específica: a ', False, '374151'),
    ('conversão automática de imagem de fórmula para código LaTeX com um único clique', True, '1E3A5F'),
    ('. Para isso é necessária uma chave de API de inteligência artificial que permita à extensão processar as imagens das fórmulas.', False, '374151'),
])

subh(doc, 'Como funciona')
numbered_list(doc, [
    'O colaborador identifica a imagem da fórmula no material (página web, PDF no navegador ou material do Canvas)',
    'Clica com o botão direito sobre a imagem',
    'Seleciona "Converter fórmula para LaTeX" no menu de contexto',
    'A extensão envia a imagem automaticamente para a API de IA configurada',
    'O modelo analisa a imagem e identifica a fórmula matemática',
    'O código LaTeX é copiado diretamente para a área de transferência — Ctrl+V já está pronto',
    'No editor LaTeX do Canvas, basta colar — a fórmula aparece renderizada corretamente',
])
sp(doc, 8)
alert(doc, '⚡', 'Tempo total: aproximadamente 5 a 10 segundos.',
    'Sem abrir nenhuma aba adicional, sem copiar para outro site, sem aguardar interfaces carregarem.')

subh(doc, 'Tecnologia utilizada')
body(doc, [
    ('A extensão utiliza o modelo ', False, '374151'),
    ('GPT-4o (OpenAI)', True, '1E3A5F'),
    (', com capacidade de visão computacional — analisa imagens e interpreta fórmulas matemáticas com alta precisão.', False, '374151'),
])
data_table(doc,
    ['Tipo de chave', 'Prefixo', 'Endpoint utilizado'],
    [
        ['OpenAI (direta)', 'sk-…',    'api.openai.com'],
        ['OpenRouter',      'sk-or-…', 'openrouter.ai'],
    ]
)
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 5. BENEFÍCIOS
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 5, 'Benefícios')
cards(doc, [
    ('⚡', 'Produtividade',           'Redução de 80 a 95% no tempo por fórmula. Em um material com 10 fórmulas, economia de 40 a 120 minutos.'),
    ('♿', 'Acessibilidade garantida', 'Toda fórmula em LaTeX é legível por leitores de tela, cumprindo as diretrizes WCAG 2.1 nível AA.'),
    ('♾', 'Sem limite de uso',        'Diferente do ChatGPT gratuito, a API é ilimitada — cada requisição é cobrada individualmente.'),
    ('🎯', 'Alta precisão',           'GPT-4o interpreta frações, somatórios, integrais, vetores e expressões complexas com alta fidelidade.'),
    ('🔗', 'Integrado ao fluxo',      'Funciona direto no navegador, sem trocar de aba ou abrir ferramentas externas.'),
    ('👥', 'Democrático',             'Não exige conhecimento em LaTeX — qualquer membro da equipe consegue usar com um clique.'),
])
subh(doc, 'Comparativo de tempo estimado')
data_table(doc,
    ['Etapa', 'Antes', 'Depois'],
    [
        ['Identificar e copiar fórmula', '~1 min',     '5 seg'],
        ['Obter código LaTeX',           '3 – 10 min', 'automático'],
        ['Colar e verificar no Canvas',  '1 – 2 min',  '~30 seg'],
    ],
    footer=['Total estimado por fórmula', '5 – 13 minutos', '~1 minuto']
)
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 6. JUSTIFICATIVA TÉCNICA
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 6, 'Justificativa Técnica da Escolha')
data_table(doc,
    ['Critério', 'ChatGPT Gratuito', 'API (OpenAI / OpenRouter)'],
    [
        ['Limite de uso',           'Sim — interrompível',       'Não — pay-as-you-go'],
        ['Integração com extensão', 'Impossível',                'Sim, via REST'],
        ['Automação (1 clique)',     'Impossível',                'Sim'],
        ['Controle da chave',       'Não',                       'Sim — revogável a qualquer momento'],
        ['Interpretação de imagem', 'Manual (usuário descreve)', 'Automática (visão computacional)'],
        ['Custo',                   'Gratuito com limitações',   'Muito baixo — centavos por fórmula'],
    ]
)
body(doc, [
    ('O GPT-4o é um modelo ', False, '374151'),
    ('multimodal', True, '1E3A5F'),
    (' — processa texto e imagens. Sua capacidade de interpretar fórmulas matemáticas a partir de imagens é superior a modelos menores ou soluções OCR tradicionais, que frequentemente falham com notações complexas.', False, '374151'),
])
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 7. CUSTOS E VIABILIDADE
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 7, 'Custos e Viabilidade')
body_p(doc, 'A API da OpenAI utiliza o modelo pay-as-you-go: paga-se apenas pelo que for usado, por token processado. Não há mensalidade, não há contrato mínimo.')

subh(doc, 'Estimativa de custo por conversão')
data_table(doc,
    ['Item', 'Quantidade estimada', 'Custo (USD)'],
    [
        ['Tokens de entrada (imagem + prompt)', '~200 – 500 tokens', '~$0,0010'],
        ['Tokens de saída (código LaTeX)',       '~50 – 200 tokens',  '~$0,0010'],
    ],
    footer=['Total por fórmula', '', '~$0,001 a $0,003']
)

# Destaque de custo
t_cost = doc.add_table(rows=1, cols=1)
tbl_no_borders(t_cost)
tbl_width_pct(t_cost)
cc = t_cost.rows[0].cells[0]
cell_shd(cc, 'DCFCE7')
cell_all_borders(cc, color='86EFAC', sz='6')
cell_margins(cc, top=160, bottom=160, left=200, right=200)
p_c = cc.paragraphs[0]
p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spc(p_c, 0, 0)
run(p_c, 'Custo médio por fórmula: ', size=12, color='166534')
run(p_c, '$0,002 USD ≈ R$ 0,01', bold=True, size=15, color='166534')
run(p_c, '  —  menos de 1 centavo por conversão', size=12, color='166534')
sp(doc, 10)

subh(doc, 'Projeção de uso mensal')
data_table(doc,
    ['Cenário', 'Fórmulas / mês', 'Custo est. (USD)', 'Custo est. (BRL)'],
    [
        ['Uso leve (1–2 FICs)',           '~50',    '~$0,10', '~R$ 0,57'],
        ['Uso moderado (5–8 FICs)',       '~300',   '~$0,60', '~R$ 3,42'],
        ['Uso intenso (FICs + revisões)', '~1.000', '~$2,00', '~R$ 11,40'],
    ]
)
alert(doc, '💡', 'Recomendação:',
    'Definir limite de gasto mensal de $5,00 USD na conta da OpenAI cobre mais de 2.000 conversões e evita cobranças inesperadas.')
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 8. PRECAUÇÕES E SEGURANÇA
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 8, 'Precauções e Segurança')

subh(doc, 'Sobre a chave de API')
bullet_list(doc, [
    'A chave é armazenada localmente no navegador — não é enviada para nenhum servidor da equipe',
    'A requisição vai diretamente do navegador para o endpoint da OpenAI/OpenRouter, sem intermediários',
    'A chave pode ser revogada a qualquer momento pelo painel da OpenAI, sem impacto a outros sistemas',
    'Recomenda-se definir um limite de gasto mensal para evitar cobranças inesperadas',
    'A chave deve ser gerada em uma conta vinculada ao responsável técnico ou à instituição',
])
sp(doc, 8)

subh(doc, 'Sobre os dados enviados')
bullet_list(doc, [
    'A extensão envia apenas a imagem da fórmula — nenhum dado do usuário ou do Canvas é transmitido',
    'Os dados de API não são usados para treinar modelos (diferente do ChatGPT gratuito)',
    'Nenhum dado de identificação pessoal é transmitido nas requisições',
])
sp(doc, 8)

subh(doc, 'Boas práticas recomendadas')
bullet_list(doc, [
    'Não compartilhar a chave em canais públicos, e-mails ou documentos não protegidos',
    'Monitorar o uso pelo painel da OpenAI ao menos uma vez por mês',
    'Revogar e substituir a chave imediatamente em caso de suspeita de vazamento',
    'Manter a extensão sempre atualizada',
])
sp(doc, 8)
alert(doc, '⚠️', 'Atenção:',
    'A chave de API possui valor financeiro — quem a tiver pode realizar requisições cobradas na conta. Trate-a como uma senha.',
    bg='FEFCE8', border='CA8A04', fc='854D0E')
hr(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 9. CONCLUSÃO
# ═══════════════════════════════════════════════════════════════════════════
section_heading(doc, 9, 'Conclusão')

t_conc = doc.add_table(rows=1, cols=1)
tbl_no_borders(t_conc)
tbl_width_pct(t_conc)
c_conc = t_conc.rows[0].cells[0]
cell_shd(c_conc, '1E3A5F')
cell_borders_none(c_conc)
cell_margins(c_conc, top=220, bottom=220, left=280, right=280)
pc1 = c_conc.paragraphs[0]
para_spc(pc1, 0, 100, line=16)
run(pc1, 'Resumo Executivo\n', bold=True, size=13, color='FFFFFF')
run(pc1, 'A adoção da chave de API para conversão automática de fórmulas LaTeX resolve um gargalo real e recorrente no fluxo de produção de materiais no Canvas LMS.', size=11, color='CBD5E1')
pc2 = c_conc.add_paragraph()
para_spc(pc2, 0, 100, line=16)
run(pc2, 'A solução proposta reduz o tempo de inserção de cada fórmula de vários minutos para menos de 10 segundos, garante acessibilidade digital e elimina a dependência de ferramentas externas com limite de uso.', size=11, color='CBD5E1')
pc3 = c_conc.add_paragraph()
para_spc(pc3, 0, 0, line=16)
run(pc3, 'O custo operacional é inferior a R$ 0,02 por fórmula. A relação custo-benefício é ', size=11, color='CBD5E1')
run(pc3, 'amplamente favorável à implementação imediata.', bold=True, size=11, color='93C5FD')

sp(doc, 16)
p_rod = doc.add_paragraph()
p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spc(p_rod, 0, 0)
run(p_rod, 'Documento interno — Equipe de Gestão do Ambiente Virtual  ·  Maio de 2026', size=9, color='94A3B8')

doc.save(r'c:\Users\Casa\servicetime-v2\relatorio-api-latex.docx')
print('Gerado com sucesso.')
